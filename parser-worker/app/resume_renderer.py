from pathlib import Path
from typing import Dict, List, Tuple

import fitz
from docx import Document


def _apply_rewrites_to_text(text: str, rewrites: List[Dict[str, str]]) -> Tuple[str, int]:
    applied = 0
    result = text
    for rewrite in rewrites:
        before = (rewrite.get("before_text") or "").strip()
        after = (rewrite.get("after_text") or "").strip()
        if not before or not after or before == after:
            continue
        if before in result:
            result = result.replace(before, after, 1)
            applied += 1
    return result, applied


def _write_plain_docx(text: str, output_path: str) -> None:
    doc = Document()
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if len(line) <= 30 and not line.endswith(("。", ".", "；", ";")):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)
    doc.save(output_path)


def _write_plain_pdf(text: str, output_path: str) -> None:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    rect = fitz.Rect(48, 48, 547, 794)
    cursor_text = text.strip() or "Resume"
    page.insert_textbox(rect, cursor_text, fontsize=10.5, fontname="helv", lineheight=1.25)
    doc.save(output_path)
    doc.close()


def render_tailored_resume(
    source_path: str,
    source_type: str,
    output_dir: str,
    base_name: str,
    rewrites: List[Dict[str, str]],
    fallback_text: str,
) -> Dict:
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_base = "".join(c for c in base_name if c.isalnum() or c in ("-", "_")) or "tailored_resume"
    docx_path = str(out_dir / f"{safe_base}.docx")
    pdf_path = str(out_dir / f"{safe_base}.pdf")
    warnings: List[str] = []
    applied_count = 0

    if source_type == "docx":
        try:
            doc = Document(source_path)
            for para in doc.paragraphs:
                original = para.text.strip()
                if not original:
                    continue
                for rewrite in rewrites:
                    before = (rewrite.get("before_text") or "").strip()
                    after = (rewrite.get("after_text") or "").strip()
                    if before and after and before in para.text and before != after:
                        para.text = para.text.replace(before, after, 1)
                        applied_count += 1
                        break
            doc.save(docx_path)
        except Exception as exc:
            warnings.append(f"DOCX original edit failed: {exc}")
            tailored_text, applied_count = _apply_rewrites_to_text(fallback_text, rewrites)
            _write_plain_docx(tailored_text, docx_path)
    else:
        tailored_text, applied_count = _apply_rewrites_to_text(fallback_text, rewrites)
        _write_plain_docx(tailored_text, docx_path)

    try:
        # Keep PDF generation deterministic and dependency-light. This PDF is
        # text-first; layout-perfect conversion can be added later via LibreOffice.
        final_text = fallback_text
        if source_type == "docx":
            try:
                final_text = "\n".join(p.text for p in Document(docx_path).paragraphs if p.text.strip())
            except Exception:
                pass
        final_text, _ = _apply_rewrites_to_text(final_text, rewrites)
        _write_plain_pdf(final_text, pdf_path)
    except Exception as exc:
        warnings.append(f"PDF render failed: {exc}")

    return {
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "applied_count": applied_count,
        "warnings": warnings,
    }
