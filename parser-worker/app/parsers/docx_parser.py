# parser-worker/app/parsers/docx_parser.py
from docx import Document
from typing import List, Dict


def parse_docx(file_path: str) -> Dict:
    doc = Document(file_path)
    full_text = ""
    paragraphs: List[str] = []
    tables: List[Dict] = []
    warnings: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
            full_text += para.text + "\n"

    for i, table in enumerate(doc.tables):
        try:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append({
                "rows": len(data),
                "cols": len(data[0]) if data else 0,
                "headers": data[0] if data else None,
                "data": data[1:] if len(data) > 1 else [],
            })
        except Exception as e:
            warnings.append(f"Table {i}: {str(e)}")

    return {
        "full_text": full_text.strip(),
        "pages": [{"text": "\n".join(paragraphs), "tables": tables, "extraction_method": "text", "confidence": 1.0}],
        "warnings": warnings,
    }
